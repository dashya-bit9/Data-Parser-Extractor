import os
import pandas as pd
from docx import Document
from docx.oxml.ns import qn


# ─── Table Extractor ──────────────────────────────────────────────────────────

def extract_tables(doc, logger):
    """
    Extract all tables from a Word document.
    Returns a list of DataFrames.
    """
    dataframes = []

    if not doc.tables:
        logger.info("  No tables found in document.")
        return dataframes

    for i, table in enumerate(doc.tables):
        data = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            data.append(cells)

        if len(data) < 2:
            continue

        headers = data[0]
        rows    = data[1:]

        # Handle duplicate or blank column names
        seen         = {}
        clean_headers = []
        for h in headers:
            h = h if h else "unnamed"
            seen[h] = seen.get(h, 0) + 1
            clean_headers.append(f"{h}_{seen[h]}" if seen[h] > 1 else h)

        df = pd.DataFrame(rows, columns=clean_headers)

        # Drop completely empty rows
        df = df[~(df == "").all(axis=1)]
        df.insert(0, "table_num", i + 1)

        dataframes.append(df)
        logger.info(f"  Table {i + 1}: {len(df)} row(s), {len(df.columns)} column(s).")

    return dataframes


# ─── Paragraph Extractor ──────────────────────────────────────────────────────

def extract_paragraphs(doc, logger):
    """
    Extract all paragraph text from a Word document.
    Includes heading level detection.
    Returns a list of dicts: {style, text}
    """
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style else "Normal"
        paragraphs.append({
            "style": style,
            "text":  text
        })

    if not paragraphs:
        logger.info("  No paragraph text found in document.")
    else:
        logger.info(f"  Extracted {len(paragraphs)} paragraph(s).")

    return paragraphs


# ─── List Extractor ───────────────────────────────────────────────────────────

def extract_lists(doc, logger):
    """
    Extract bullet and numbered list items from a Word document.
    Returns a list of dicts: {type, text}
    """
    items = []

    for para in doc.paragraphs:
        text  = para.text.strip()
        style = para.style.name if para.style else ""

        if not text:
            continue

        if "List" in style or "Bullet" in style:
            list_type = "bullet" if "Bullet" in style else "numbered"
            items.append({
                "type": list_type,
                "text": text
            })

    if not items:
        logger.info("  No list items found in document.")
    else:
        logger.info(f"  Extracted {len(items)} list item(s).")

    return items


# ─── Metadata Extractor ───────────────────────────────────────────────────────

def extract_metadata(doc, logger):
    """
    Extract core document metadata (author, title, dates, etc).
    Returns a dict.
    """
    props = doc.core_properties
    metadata = {
        "title":          props.title    or "",
        "author":         props.author   or "",
        "created":        str(props.created)  if props.created  else "",
        "modified":       str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
        "subject":        props.subject  or "",
        "description":    props.description or "",
        "keywords":       props.keywords or "",
    }

    logger.info(f"  Metadata — Title: '{metadata['title']}' | Author: '{metadata['author']}'")
    return metadata


# ─── Empty Row Filter ─────────────────────────────────────────────────────────

def drop_empty_rows(df, logger):
    """Drop rows where every cell is empty or None."""
    before = len(df)
    df     = df.dropna(how="all")
    df     = df[~(df == "").all(axis=1)]
    after  = len(df)

    if before != after:
        logger.info(f"  Dropped {before - after} empty row(s).")

    return df


# ─── Main Entry Point ─────────────────────────────────────────────────────────

def extract(filepath, logger):
    """
    Main Word document entry point.
    Tries tables first, falls back to paragraphs.
    Returns a dict matching the same format as pdf_extractor.extract():
        - mode: 'table' or 'text'
        - data: list of DataFrames (table) or list of dicts (text)
        - page_count: number of tables or paragraphs found
        - source: filename
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Document not found: {filepath}")

    filename = os.path.basename(filepath)
    logger.info(f"Opening Word document: {filename}")

    doc      = Document(filepath)
    metadata = extract_metadata(doc, logger)
    tables   = extract_tables(doc, logger)

    # ── Table mode ──
    if tables:
        return {
            "mode":       "table",
            "data":       tables,
            "page_count": len(tables),
            "source":     filename,
            "metadata":   metadata
        }

    # ── Text fallback ──
    logger.info("No tables found. Falling back to paragraph extraction...")
    paragraphs = extract_paragraphs(doc, logger)

    if paragraphs:
        return {
            "mode":       "text",
            "data":       paragraphs,
            "page_count": len(paragraphs),
            "source":     filename,
            "metadata":   metadata
        }

    # ── List fallback ──
    logger.info("No paragraphs found. Falling back to list extraction...")
    lists = extract_lists(doc, logger)

    return {
        "mode":       "text",
        "data":       lists,
        "page_count": len(lists),
        "source":     filename,
        "metadata":   metadata
    }